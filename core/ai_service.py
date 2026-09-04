try:
    from google import genai
    from google.genai import types
    HAS_GENAI = True
except ImportError:
    HAS_GENAI = False

from PIL import Image
import io

from django.conf import settings
import logging
import os

logger = logging.getLogger("pos.ai")

# One place to change which Gemini model every AI feature uses -- was
# hardcoded as a literal string in four separate call sites, meaning
# switching models (e.g. pinning a specific version instead of "-latest",
# or trying a cheaper/newer one later) needed a code change in each. Now
# it's an env var, defaulting to the same "-latest" alias as before, so
# nothing changes unless GEMINI_MODEL_NAME is actually set.
GEMINI_MODEL_NAME = os.getenv("GEMINI_MODEL_NAME", "gemini-flash-latest")


class AIService:
    def __init__(self):
        import os
        self.api_key = os.getenv("GOOGLE_API_KEY")
        self.client = None

        # ══════════════════════════════════════════════════════════════════
        # TEMPORARY DEV-STAGE WORKAROUND — NOT A PERMANENT DESIGN.
        #
        # A second Google account's key, tried only when the primary key's
        # call actually fails at runtime (free-tier quota exhausted, rate
        # limited, etc). This exists purely because Rasova has no paying
        # customers yet, so every request today is our own testing burning
        # through a 20/day free-tier cap — not because "two free accounts"
        # is the right long-term architecture. Relying on multiple free-tier
        # accounts to route around a rate limit is exactly what Google's
        # terms discourage; keep this short-lived.
        #
        # >>> REMINDER: the day Rasova gets its first real paying customer,
        # >>> retire this. Upgrade GOOGLE_API_KEY's account to a paid Gemini
        # >>> tier (removes the cap entirely, costs pennies at real volume),
        # >>> then delete GOOGLE_API_KEY_FALLBACK from .env. Don't let this
        # >>> quietly become permanent just because it works.
        #
        # Full writeup: md_files/ai_cost_plan_eli5_2026-09-04.html (step 5).
        # ══════════════════════════════════════════════════════════════════
        self.fallback_api_key = os.getenv("GOOGLE_API_KEY_FALLBACK")
        self.fallback_client = None

        # Log the SPECIFIC reason the client can't be built, so the server logs
        # say exactly what to fix instead of a generic "no AI key". The #1 cause
        # in practice is a stale process: .env was updated but the worker that
        # runs the import (celery for image import, gunicorn for the sync
        # fallback) was never restarted, so it still has the old environment.
        if not HAS_GENAI:
            logger.warning(
                "AI disabled: the 'google-genai' package is not installed in this "
                "environment. Run: pip install -r requirements.txt"
            )
            return
        if not self.api_key:
            logger.warning(
                "AI disabled: GOOGLE_API_KEY is not set in THIS process's environment. "
                "Set it in .env, then RESTART the service that runs the import "
                "(celery for image import, gunicorn for the sync fallback)."
            )
            return
        try:
            self.client = genai.Client(api_key=self.api_key)
            logger.info("AI client initialized (key ...%s).", self.api_key[-4:])
        except Exception as e:
            logger.error("Failed to initialize AI client — check the key is valid: %s", e)

        if self.fallback_api_key:
            try:
                self.fallback_client = genai.Client(api_key=self.fallback_api_key)
                logger.info("AI fallback client initialized (key ...%s) -- dev-stage only, see __init__.", self.fallback_api_key[-4:])
            except Exception as e:
                logger.error("Failed to initialize AI fallback client: %s", e)

    def _generate_content(self, model, contents):
        """Every Gemini call in this file goes through here instead of hitting
        self.client directly, so the fallback-account retry lives in exactly
        one place rather than duplicated at each call site. Tries the primary
        key; only if that call itself fails (quota, rate limit, transient
        error) AND a second key is configured does it retry once on the
        second key. If neither fallback_client exists nor the retry helps,
        the original exception propagates exactly as it always did, so every
        existing except-block (regex-parser fallback, error messages) keeps
        working unchanged."""
        try:
            return self.client.models.generate_content(model=model, contents=contents)
        except Exception as e:
            if not self.fallback_client:
                raise
            logger.warning("Primary Gemini key failed (%s) -- retrying on fallback key.", e)
            return self.fallback_client.models.generate_content(model=model, contents=contents)

    def _resize_image(self, image_bytes, max_size=(1024, 1024)):
        """Re-encode any image to a clean, right-sized RGB JPEG.

        Raises if the bytes aren't a readable image (e.g. an iPhone HEIC, or a
        corrupt upload). The old version swallowed that and returned the ORIGINAL
        bytes while still labelling them image/jpeg — so Gemini received
        non-JPEG data claiming to be JPEG and replied with the opaque
        "Unable to process input image" 400. Better to fail here and let the
        caller show a "use a JPG/PNG" message.
        """
        from PIL import ImageOps
        img = Image.open(io.BytesIO(image_bytes))
        img = ImageOps.exif_transpose(img)          # honor phone-camera rotation
        if img.mode != "RGB":
            img = img.convert("RGB")                # JPEG needs RGB (handles P/RGBA/CMYK/LA/…)
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        output = io.BytesIO()
        img.save(output, format="JPEG", quality=80, optimize=True)
        return output.getvalue()

    # -------------------------------------------------------
    # FILE ROUTING — turn any uploaded file into something Gemini can read
    # -------------------------------------------------------

    def _file_to_content(self, file_bytes, mime_type):
        """Return a Gemini content item for an uploaded menu file.

        - Images  -> re-encoded JPEG Part
        - PDFs    -> PDF Part (Gemini reads text AND scanned/image PDFs natively)
        - Word    -> extracted text
        - Excel   -> extracted text
        - text/csv-> decoded text
        Raises a friendly error for anything unreadable.
        """
        mt = (mime_type or "").lower()

        # PDF — sniff the magic bytes too, in case the browser sent no mime type
        if mt == "application/pdf" or file_bytes[:5] == b"%PDF-":
            return types.Part.from_bytes(data=file_bytes, mime_type="application/pdf")

        # Word (.docx / .doc)
        if "wordprocessingml" in mt or mt == "application/msword":
            return "Menu text extracted from a Word document:\n" + self._extract_docx(file_bytes)

        # Excel (.xlsx / .xls)
        if "spreadsheetml" in mt or "ms-excel" in mt:
            return "Menu text extracted from a spreadsheet:\n" + self._extract_xlsx(file_bytes)

        # Plain text / CSV
        if mt.startswith("text/"):
            return "Menu text:\n" + file_bytes.decode("utf-8", errors="replace")

        # Otherwise assume it's a photo (explicit image/* or unknown type).
        try:
            return types.Part.from_bytes(
                data=self._resize_image(file_bytes), mime_type="image/jpeg"
            )
        except Exception as e:
            logger.warning("Unreadable upload (mime=%s): %s", mime_type, e)
            raise Exception(
                "Couldn't read that file. Please upload a photo (JPG/PNG), a PDF, "
                "a Word document (.docx), an Excel sheet (.xlsx), or plain text/CSV. "
                "(iPhone HEIC photos aren't supported — screenshot the menu instead.)"
            )

    def _extract_docx(self, file_bytes):
        """Pull all paragraph + table text out of a .docx."""
        try:
            from docx import Document
        except ImportError:
            raise Exception(
                "Word import isn't available on the server yet. Save the menu as "
                "a PDF or take a screenshot, and upload that instead."
            )
        doc = Document(io.BytesIO(file_bytes))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append("  ".join(cells))
        return "\n".join(lines)

    def _extract_pdf_text(self, file_bytes):
        """Pull embedded text out of a PDF, page by page. Only works for a
        genuine digital PDF (text layer present) -- a scanned/photographed
        menu saved as PDF has no text layer, and this correctly comes back
        empty rather than guessing. Used only as the fallback source when
        Gemini itself (which reads scanned PDFs too, via OCR) is unavailable;
        the normal working path still sends the raw PDF straight to Gemini,
        since that's strictly more capable than this when it's up."""
        try:
            from pypdf import PdfReader
        except ImportError:
            return ""
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            logger.warning("PDF text extraction failed: %s", e)
            return ""

    # Generous headroom over any real menu (a huge one might run a few
    # thousand characters) while still bounding worst case -- this is the
    # backstop against a decompression-bomb-style file: something whose
    # compressed size passes the 15MB upload cap but unpacks internally
    # (inside pypdf/python-docx/openpyxl) into far more text than that.
    # Truncating here means the regex parser downstream never sees more
    # than this, regardless of how much the file expanded to internally.
    _MAX_FALLBACK_TEXT_CHARS = 50_000

    def _extract_fallback_text(self, file_bytes, mime_type):
        """Best-effort plain text for the no-AI fallback parser, from
        whatever file type was uploaded. Mirrors _file_to_content's routing,
        but every branch here degrades to "" instead of raising -- this is
        already the fallback path, so there's nothing lower to fall back to.
        Images return "" deliberately: the regex parser needs text, and
        there's no local OCR to get any from a photo."""
        mt = (mime_type or "").lower()
        if mt == "application/pdf" or file_bytes[:5] == b"%PDF-":
            result = self._extract_pdf_text(file_bytes)
        elif "wordprocessingml" in mt or mt == "application/msword":
            try:
                result = self._extract_docx(file_bytes)
            except Exception:
                result = ""
        elif "spreadsheetml" in mt or "ms-excel" in mt:
            try:
                result = self._extract_xlsx(file_bytes)
            except Exception:
                result = ""
        elif mt.startswith("text/"):
            result = file_bytes.decode("utf-8", errors="replace")
        else:
            result = ""
        return result[: self._MAX_FALLBACK_TEXT_CHARS]

    def _extract_xlsx(self, file_bytes):
        """Flatten every non-empty cell of every sheet into text (openpyxl is
        already a dependency for the GSTR-1 export)."""
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        lines = []
        try:
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append("  ".join(cells))
        finally:
            wb.close()
        return "\n".join(lines)

    def parse_menu(self, text=None, image_bytes=None, mime_type=None):
        """
        Parses menu from text or image. 
        Falls back to manual regex parsing if AI is not configured and only text is provided.
        """
        # Whatever plain text we can get our hands on without needing Gemini at
        # all -- either what was pasted directly, or (for Word/Excel/CSV/PDF
        # uploads) text extracted locally from the file. This is the one thing
        # the regex fallback parser can work with; a bare photo has no local
        # text to extract, so this stays "" for that case.
        fallback_text = text.strip() if text and text.strip() else ""
        if not fallback_text and image_bytes:
            fallback_text = self._extract_fallback_text(image_bytes, mime_type)

        if not self.client:
            if fallback_text:
                logger.info("Google API Key not found. Falling back to manual text parser.")
                return self._manual_text_parse(fallback_text)

            # A real photo genuinely needs AI/OCR -- no local text to fall back to.
            if image_bytes:
                raise Exception("Image parsing requires an AI activation key. Please add it to your configuration.")

            # If they provided neither or empty text
            raise Exception("No menu data provided to parse.")

        prompt = """
        Analyze this restaurant menu and extract all items.
        Format your response as a valid JSON list of objects. Each object represents a category.
        Each category object must have:
        1. "category": String (name of the category, e.g., "Starters")
        2. "items": List of objects, each with:
           - "name" (String)
           - "price" (Float/Number)
           - "is_veg" (Boolean): true if the dish is vegetarian (no meat, poultry,
             fish, seafood, or egg), false otherwise. Judge from the dish name and
             any veg/non-veg marker shown on the menu (a green or red dot/square is
             the standard Indian menu convention: green = veg, red/brown = non-veg).
             If genuinely ambiguous, default to false rather than falsely labeling
             a non-veg dish as vegetarian.

        If you find a description, ignore it. Only extract name, price, and is_veg.
        If no price is found, use 0.
        Output ONLY the raw JSON list. No markdown, no backticks.
        """

        contents = [prompt]
        if text and text.strip():
            contents.append(f"Here is the menu text:\n{text}")
        
        if image_bytes:
            # `image_bytes` is really "the uploaded file" — could be an image, a
            # PDF, a Word doc, an Excel sheet or plain text. _file_to_content
            # routes each to the right thing (a Gemini Part for images/PDFs, or
            # extracted text for Office docs).
            contents.append(self._file_to_content(image_bytes, mime_type))

        try:
            response = self._generate_content(GEMINI_MODEL_NAME, contents)
            import json

            # The new SDK response structure: response.text or response.candidates[0].content.parts[0].text
            res_text = response.text or ""
            # Strip potential markdown backticks
            raw_json = res_text.strip().replace("```json", "").replace("```", "")
            return json.loads(raw_json)
        except Exception as e:
            logger.error("In-build AI API Error: %s", e)
            if fallback_text:
                logger.info("Falling back to manual parser after Gemini API error.")
                return self._manual_text_parse(fallback_text)
            raise e

    def _manual_text_parse(self, text):
        """
        Fallback parser using regex to extract categories and items from raw text.
        """
        import re
        if not text: return []

        lines = [line.strip() for line in text.split('\n') if line.strip()]

        # A flattened table export (found via a real menu PDF during testing:
        # pypdf turns each "Category | Dish | Price" table row into one line,
        # "Category Dish Price", not a separate category header followed by
        # item lines). Only attempt to split an inline category off the front
        # of each item name when there's clear positive evidence this file is
        # that shape -- a column-header row like "Category Dish Price" -- not
        # just because some leading word happens to repeat, which would
        # misfire on an ordinary menu with e.g. "Chicken Biryani" and
        # "Chicken 65" both starting with "Chicken".
        header_line = next((l for l in lines if self._looks_like_table_header(l)), None)
        inline_categories = self._detect_inline_categories(lines) if header_line else set()

        structured = []
        structured_by_name = {}
        current_category = {"category": "General", "items": []}

        def _target_category(name):
            cat_name, item_name = self._split_inline_category(name, inline_categories)
            if not cat_name:
                return current_category, name
            cat = structured_by_name.get(cat_name)
            if not cat:
                cat = {"category": cat_name, "items": []}
                structured_by_name[cat_name] = cat
                structured.append(cat)
            return cat, item_name

        for line in lines:
            if line is header_line:
                continue  # the column-header row itself is not a category

            # Detect category: All caps, or ends with colon, or short line with no digits
            is_cat = (line.isupper() and len(line) > 3) or line.endswith(':') or (not any(c.isdigit() for c in line) and len(line) < 25)

            if is_cat and not inline_categories:
                if current_category["items"]:
                    structured.append(current_category)
                current_category = {"category": line.rstrip(':').strip().title(), "items": []}
                continue

            # Match "Item Name 123.45" or "Item Name - 123". Deliberately not
            # a single regex like r'(.*?)(?:[:\-\s]+)(\d+...)$' -- that shape
            # lets the non-greedy name group and the separator class both
            # match '-'/whitespace, so the engine can split them ambiguously
            # in exponentially many ways on a pathological line (e.g. a long
            # run of dashes/spaces with no trailing digits), which is exactly
            # the ReDoS CodeQL flagged. A plain trailing-digits match has no
            # such ambiguity, so peel the price off with that, then take
            # whatever's left as the name with ordinary string slicing.
            price_match = re.search(r'(\d+(?:\.\d+)?)$', line)
            if price_match:
                raw_name = line[:price_match.start()].strip(' \t:-')
                target, name = _target_category(raw_name)
                target["items"].append({
                    "name": name,
                    "price": float(price_match.group(1)),
                    "is_veg": self._guess_veg(name),
                })
            else:
                # No price found, add as item with price 0
                raw_name = line.strip()
                target, name = _target_category(raw_name)
                target["items"].append({
                    "name": name,
                    "price": 0,
                    "is_veg": self._guess_veg(name),
                })

        if current_category["items"]:
            structured.append(current_category)
        return structured

    _TABLE_HEADER_WORDS = ("category", "dish", "item", "price", "name", "menu")

    def _looks_like_table_header(self, line):
        """A column-header row from a table export ("Category Dish Price"),
        not real menu content -- recognized by containing 2+ generic column
        labels and having no price of its own. This is the positive-evidence
        gate for the inline-category splitting below: without it, a plain
        menu with no such header row is left completely untouched."""
        import re
        if re.search(r'\d', line):
            return False
        lname = line.lower()
        return sum(1 for w in self._TABLE_HEADER_WORDS if w in lname) >= 2

    def _detect_inline_categories(self, lines):
        """Only called once a table-header row has already confirmed this
        file is a flattened table export. Finds the leading word-phrase each
        item's category got flattened into -- the phrase that recurs across
        multiple item lines (e.g. "Main Course" appearing at the start of
        five different item lines). Checked at 1, 2, and 3 leading words so
        _split_inline_category can prefer the longest match per line."""
        import re
        from collections import Counter
        candidates = []
        for line in lines:
            m = re.search(r'(\d+(?:\.\d+)?)$', line)
            if m:
                name = line[:m.start()].strip(' \t:-')
                if name:
                    candidates.append(name.split())

        recurring = set()
        for n in (1, 2, 3):
            counts = Counter(" ".join(words[:n]) for words in candidates if len(words) > n)
            recurring |= {phrase for phrase, count in counts.items() if count >= 2}
        return recurring

    def _split_inline_category(self, name, inline_categories):
        """If `name` starts with a detected inline category phrase, split it
        off and return (category, remaining item name) -- preferring the
        longest matching prefix ("Main Course" over just "Main"). Returns
        (None, name) unchanged when nothing matches, which is always true
        when inline_categories is empty (the ordinary, non-tabular case)."""
        if not inline_categories:
            return None, name
        words = name.split()
        for n in (3, 2, 1):
            if len(words) > n:
                prefix = " ".join(words[:n])
                if prefix in inline_categories:
                    return prefix.title(), " ".join(words[n:])
        return None, name

    # Non-veg keywords only -- deliberately not the reverse (a "contains veg
    # keyword" allowlist), since a dish can be non-veg without ever naming an
    # ingredient (e.g. "Chef's Special"). Defaulting unmatched names to veg
    # is still a guess, but it's the same default the model field itself
    # already uses, so this only removes the *confidently wrong* cases.
    _NON_VEG_KEYWORDS = (
        "chicken", "mutton", "lamb", "beef", "pork", "bacon", "ham",
        "fish", "prawn", "shrimp", "crab", "squid", "calamari", "seafood",
        "egg", "keema", "meat",
    )

    def _guess_veg(self, name):
        """Best-effort veg/non-veg guess for the regex fallback path, which
        has no real language understanding. Keyword-based, not a substitute
        for the AI classification in parse_menu -- used only when no API key
        is configured. Deliberately only matches actual protein/ingredient
        names, not cooking styles like "tikka" or "kebab" -- those appear on
        plenty of real vegetarian dishes (Paneer Tikka), so including them
        would create confident-but-wrong false positives, which is worse
        than the plain "default to veg" the model field already does."""
        lname = name.lower()
        if "veg" in lname and "non veg" not in lname and "non-veg" not in lname:
            return True
        return not any(kw in lname for kw in self._NON_VEG_KEYWORDS)

    def parse_recipe(self, text=None, image_bytes=None, mime_type=None):
        """
        Extracts a flat ingredient list from a recipe document/photo: each
        line is the ingredient name and the quantity exactly as written
        ("1 cup", "a pinch", "200g") — no unit conversion, no inventory
        matching. Those are separate, independently-debuggable steps done by
        inventory/recipe_unit_table.py and match_ingredients() below, so a
        bad extraction and a bad match/conversion never get tangled together
        in one opaque failure.

        Unlike parse_menu, there is no manual-regex fallback — a recipe's
        prose/format varies far more than a price list, so without AI this
        just raises a clear "AI not configured" error rather than guessing.
        """
        if not self.client:
            raise Exception(
                "AI recipe import requires an AI activation key. Please add it to your configuration."
            )
        if not (text and text.strip()) and not image_bytes:
            raise Exception("No recipe document provided to parse.")

        prompt = """
        You are extracting the ingredient list from a recipe document for ONE dish.
        The document may be messy: handwritten, split into sections (e.g. "For the
        marinade:", "For the gravy:"), or written as prose rather than a clean list.
        Read the whole thing and pull out every ingredient, from every section.

        Ignore preparation steps, cooking instructions, serving suggestions, and
        anything that isn't an ingredient with a quantity.

        Format your response as a valid JSON list of objects, one per ingredient:
        1. "ingredient": String — the ingredient name only (e.g. "Onion", not "2 medium onions")
        2. "quantity_text": String — the quantity EXACTLY as written in the source,
           unmodified (e.g. "1 cup", "2 medium", "a pinch", "200g", "1/2 tsp").
           Do not convert, round, or normalize it — copy it verbatim.

        If the same ingredient appears more than once (e.g. once in a marinade and
        once as a garnish), output it as two separate entries rather than combining them.

        Output ONLY the raw JSON list. No markdown, no backticks, no commentary.
        """

        contents = [prompt]
        if text and text.strip():
            contents.append(f"Here is the recipe text:\n{text}")
        if image_bytes:
            contents.append(self._file_to_content(image_bytes, mime_type))

        import json
        try:
            response = self._generate_content(GEMINI_MODEL_NAME, contents)
            res_text = response.text or ""
            raw_json = res_text.strip().replace("```json", "").replace("```", "")
            parsed = json.loads(raw_json)
            if not isinstance(parsed, list):
                raise ValueError("Expected a JSON list of ingredients")
            return parsed
        except Exception as e:
            logger.error("AI recipe extraction error: %s", e)
            raise Exception("Couldn't read that recipe. Try a clearer photo or a text/PDF export.")

    def match_ingredients(self, unmatched_names, inventory_names):
        """
        One batched call resolving every ingredient name rapidfuzz couldn't
        confidently match, against this tenant+outlet's real inventory item
        names. Deliberately a single call for the whole leftover list (not
        one call per ingredient) to keep cost/latency bounded.

        Returns {ingredient_name: matched_inventory_name_or_None}. Any name
        missing from the response (bad JSON, Gemini omitted it, etc.) is
        treated as unmatched by the caller — never assumed matched.
        """
        import json

        if not unmatched_names:
            return {}
        if not self.client or not inventory_names:
            return {name: None for name in unmatched_names}

        prompt = f"""
        You are matching recipe ingredient names to a restaurant's existing
        inventory item names. For each ingredient below, pick the inventory
        item name that refers to the SAME real-world ingredient (matching
        synonyms/variants is fine, e.g. "cream" -> "Heavy Cream"), or null if
        none of the inventory items are actually the same ingredient. Do not
        pick a loosely related item just because no better option exists —
        null is the correct answer when unsure.

        Ingredients to match: {json.dumps(unmatched_names)}

        Inventory item names to match against: {json.dumps(inventory_names)}

        Respond with ONLY a raw JSON object mapping each ingredient name (exactly
        as given above) to either a matching inventory item name (exactly as
        given above) or null. No markdown, no backticks, no commentary.
        """

        try:
            response = self._generate_content(GEMINI_MODEL_NAME, [prompt])
            res_text = response.text or ""
            raw_json = res_text.strip().replace("```json", "").replace("```", "")
            parsed = json.loads(raw_json)
            if not isinstance(parsed, dict):
                raise ValueError("Expected a JSON object")
            # Only trust names that were actually offered, matched to names
            # that actually exist in inventory — anything else falls back to
            # unmatched rather than trusting an AI-invented name.
            inventory_set = set(inventory_names)
            result = {}
            for name in unmatched_names:
                match = parsed.get(name)
                result[name] = match if match in inventory_set else None
            return result
        except Exception as e:
            logger.error("AI ingredient matching error: %s", e)
            return {name: None for name in unmatched_names}

    def suggest_pricing(self, dish_name, ingredients_with_costs):
        """
        Suggests pricing for a dish based on ingredient costs and 30% food cost target.
        """
        if not self.client: return None
        
        prompt = f"Calculate the suggested menu price for '{dish_name}'. Ingredients and their total cost in this dish: {ingredients_with_costs}. Aim for a 30% food cost percentage. Add a summary of why."
        try:
            response = self._generate_content(GEMINI_MODEL_NAME, prompt)
            return response.text
        except Exception as e:
            logger.error("AI pricing suggestion error: %s", e)
            return None
